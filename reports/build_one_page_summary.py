from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


HERE = Path(__file__).resolve().parent
DOCX_PATH = HERE / "TrustLens_AI_One_Page_Research_Summary.docx"
PDF_PATH = HERE / "TrustLens_AI_One_Page_Research_Summary.pdf"

NAVY = "17365D"
BLUE = "2E74B5"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GRAY = "5B6573"
INK = "20252B"


def font(run, *, size=10.2, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def cell_margins(cell, top=70, bottom=70, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_margins(row.cells[index])


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=11.3, bold=True, color=BLUE)
    return p


def add_body(doc, text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    font(p.add_run(text))
    return p


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.text = "TRUSTLENS AI  |  SUPERVISOR RESEARCH BRIEF"
    for run in header.runs:
        font(run, size=8.2, bold=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    font(p.add_run("TRUSTLENS AI"), size=23, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run("Human-Governed Machine Learning for Reliability-Aware Risk Classification"), size=12.7, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    font(p.add_run("Akshith Moharampudi | MComp Computer Science | Responsible AI and Machine Learning"), size=9.6, color=GRAY)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [7.0])
    shade(callout.cell(0, 0), LIGHT)
    cp = callout.cell(0, 0).paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    font(cp.add_run("Research question. "), bold=True, color=NAVY)
    font(cp.add_run("Can calibrated prediction, drift and out-of-distribution checks, explainability and explicit human-review rules expose unreliable model outputs more effectively than confidence scores alone?"))

    add_heading(doc, "Why this project matters")
    add_body(doc, "High-stakes models can appear confident while operating outside their evidence base. TrustLens treats uncertainty, monitoring and escalation as executable system behaviour rather than documentation added after model development.")

    add_heading(doc, "Research design")
    add_body(doc, "The case study uses 1,000 historical South German Credit records. A stratified 20% test partition was reserved before development. Logistic regression, random forest and gradient-boosting candidates were evaluated with five-fold cross-validation, asymmetric error cost and calibration measures. The selected cost-sensitive random forest was sigmoid-calibrated, used a threshold fixed on development evidence and excluded two sensitive or ambiguous inputs. It was evaluated once on the untouched 200-record partition.")

    metrics = doc.add_table(rows=2, cols=5)
    set_table_geometry(metrics, [1.4] * 5)
    headers = ("Balanced accuracy", "Higher-risk recall", "Precision", "Weighted cost", "Automated tests")
    values = ("0.656", "0.833", "0.407", "123 vs 300", "27 passing")
    for i, value in enumerate(headers):
        metrics.cell(0, i).text = value
        shade(metrics.cell(0, i), NAVY)
        for run in metrics.cell(0, i).paragraphs[0].runs:
            font(run, size=8.5, bold=True, color="FFFFFF")
    for i, value in enumerate(values):
        metrics.cell(1, i).text = value
        shade(metrics.cell(1, i), PALE)
        metrics.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in metrics.cell(1, i).paragraphs[0].runs:
            font(run, size=10, bold=True, color=NAVY)

    add_heading(doc, "Governance and reliability evidence")
    add_body(doc, "The prototype implements probability calibration, population-drift detection, individual out-of-distribution screening, local sensitivity explanations, subgroup diagnostics and pause-or-review rules. A 20% development review budget captured 23% of observed errors. Post-hoc 95% Wilson intervals report higher-risk recall as 0.720-0.907 and precision as 0.324-0.495.")

    add_heading(doc, "Failure analysis and limitations")
    add_body(doc, "The model produced 73 false positives and 10 false negatives. Error-slice diagnostics found sharply different failure rates across recorded checking-account-status codes, reinforcing the need for human review. These results are descriptive, not causal: the dataset is small, dates from 1973-1975, contains limited demographic information and has no external validation. TrustLens is an educational research prototype and must not make real lending or eligibility decisions.")

    add_heading(doc, "Doctoral research direction")
    add_body(doc, "A doctoral extension would test governance-aware model selection on contemporary multi-institutional data, evaluate selective prediction and abstention, study realistic temporal drift, and measure whether human reviewers can correct errors without automation bias. The broader goal is responsible AI that expands access to technology while making uncertainty and system limits visible.", after=5)

    links = doc.add_table(rows=2, cols=2)
    set_table_geometry(links, [1.4, 5.6])
    rows = (
        ("Repository", "https://github.com/akshith2001/trustlens-ai"),
        ("Live dashboard", "https://trustlens-governance-ai.streamlit.app/"),
    )
    for i, (label, value) in enumerate(rows):
        links.cell(i, 0).text = label
        links.cell(i, 1).text = value
        shade(links.cell(i, 0), PALE)
        for run in links.cell(i, 0).paragraphs[0].runs:
            font(run, size=9, bold=True, color=NAVY)
        for run in links.cell(i, 1).paragraphs[0].runs:
            font(run, size=9, color=BLUE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Research portfolio summary | 22 August 2026"), size=8, color=GRAY)

    doc.core_properties.title = "TrustLens AI One-Page Research Summary"
    doc.core_properties.author = "Akshith Moharampudi"
    doc.core_properties.subject = "Responsible AI doctoral research portfolio"
    doc.save(DOCX_PATH)


def build_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=11.2, textColor=colors.HexColor("#20252B"), spaceAfter=4)
    h = ParagraphStyle("H", parent=body, fontName="Helvetica-Bold", fontSize=10.8, leading=12.5, textColor=colors.HexColor("#2E74B5"), spaceBefore=6, spaceAfter=2)
    title = ParagraphStyle("Title", parent=body, fontName="Helvetica-Bold", fontSize=22, leading=24, textColor=colors.HexColor("#17365D"), spaceAfter=1)
    subtitle = ParagraphStyle("Subtitle", parent=body, fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor("#2E74B5"), spaceAfter=2)
    muted = ParagraphStyle("Muted", parent=body, fontSize=8.8, textColor=colors.HexColor("#5B6573"), spaceAfter=6)
    call = ParagraphStyle("Call", parent=body, fontSize=9.2, leading=11.2, spaceAfter=0)
    small = ParagraphStyle("Small", parent=body, fontSize=8, leading=9.5, alignment=TA_LEFT, spaceAfter=0)
    small_header = ParagraphStyle("SmallHeader", parent=small, fontName="Helvetica-Bold", textColor=colors.white)

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=LETTER, leftMargin=.72*inch, rightMargin=.72*inch, topMargin=.5*inch, bottomMargin=.45*inch, title="TrustLens AI One-Page Research Summary", author="Akshith Moharampudi")
    story = [
        Paragraph("TRUSTLENS AI", title),
        Paragraph("Human-Governed Machine Learning for Reliability-Aware Risk Classification", subtitle),
        Paragraph("Akshith Moharampudi | MComp Computer Science | Responsible AI and Machine Learning", muted),
    ]
    callout = Table([[Paragraph("<b>Research question.</b> Can calibrated prediction, drift and out-of-distribution checks, explainability and explicit human-review rules expose unreliable model outputs more effectively than confidence scores alone?", call)]], colWidths=[7.06*inch])
    callout.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#F4F6F9")),("BOX",(0,0),(-1,-1),.6,colors.HexColor("#9FB4CA")),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6)]))
    story += [callout, Paragraph("Why this project matters", h), Paragraph("High-stakes models can appear confident while operating outside their evidence base. TrustLens treats uncertainty, monitoring and escalation as executable system behaviour rather than documentation added after model development.", body), Paragraph("Research design", h), Paragraph("The case study uses 1,000 historical South German Credit records. A stratified 20% test partition was reserved before development. Logistic regression, random forest and gradient-boosting candidates were evaluated with five-fold cross-validation, asymmetric error cost and calibration measures. The selected cost-sensitive random forest was sigmoid-calibrated, used a threshold fixed on development evidence and excluded two sensitive or ambiguous inputs. It was evaluated once on the untouched 200-record partition.", body)]

    metrics_data = [[Paragraph(x, small_header) for x in ("Balanced accuracy", "Higher-risk recall", "Precision", "Weighted cost", "Automated tests")], [Paragraph(f"<b>{x}</b>", small) for x in ("0.656", "0.833", "0.407", "123 vs 300", "27 passing")]]
    metrics = Table(metrics_data, colWidths=[1.412*inch]*5)
    metrics.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17365D")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("BACKGROUND",(0,1),(-1,1),colors.HexColor("#E8EEF5")),("ALIGN",(0,0),(-1,-1),"CENTER"),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#B9C2CE")),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story += [metrics, Paragraph("Governance and reliability evidence", h), Paragraph("The prototype implements probability calibration, population-drift detection, individual out-of-distribution screening, local sensitivity explanations, subgroup diagnostics and pause-or-review rules. A 20% development review budget captured 23% of observed errors. Post-hoc 95% Wilson intervals report higher-risk recall as 0.720-0.907 and precision as 0.324-0.495.", body), Paragraph("Failure analysis and limitations", h), Paragraph("The model produced 73 false positives and 10 false negatives. Error-slice diagnostics found sharply different failure rates across recorded checking-account-status codes, reinforcing the need for human review. These results are descriptive, not causal: the dataset is small, dates from 1973-1975, contains limited demographic information and has no external validation. TrustLens is an educational research prototype and must not make real lending or eligibility decisions.", body), Paragraph("Doctoral research direction", h), Paragraph("A doctoral extension would test governance-aware model selection on contemporary multi-institutional data, evaluate selective prediction and abstention, study realistic temporal drift, and measure whether human reviewers can correct errors without automation bias. The broader goal is responsible AI that expands access to technology while making uncertainty and system limits visible.", body)]

    links_data = [[Paragraph("<b>Repository</b>", small), Paragraph("https://github.com/akshith2001/trustlens-ai", small)], [Paragraph("<b>Live dashboard</b>", small), Paragraph("https://trustlens-governance-ai.streamlit.app/", small)]]
    links = Table(links_data, colWidths=[1.3*inch,5.76*inch])
    links.setStyle(TableStyle([("BACKGROUND",(0,0),(0,-1),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#B9C2CE")),("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story += [links, Spacer(1,4), Paragraph("Research portfolio summary | 22 August 2026", muted)]
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
