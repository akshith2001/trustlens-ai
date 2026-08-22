from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent
DOCX_PATH = OUT / "TrustLens_AI_Research_Report.docx"
FIGURE = ROOT / "figures" / "final_test_summary.png"

NAVY = "17365D"
BLUE = "2E74B5"
PALE = "EAF1F8"
LIGHT = "F3F5F7"
GRAY = "5B6573"
WHITE = "FFFFFF"
RED = "9B1C1C"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header_and_style(table, widths):
    set_table_widths(table, widths)
    set_repeat_header(table.rows[0])
    for cell in table.rows[0].cells:
        shade(cell, NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9.5)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", *, bold_lead=None, italic=False, align=None, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label + " ")
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.45)
sec.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in (
    ("Title", 27, 0, 8, NAVY),
    ("Subtitle", 13, 0, 8, GRAY),
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 11.5, 9, 4, NAVY),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.2

header = sec.header
hp = header.paragraphs[0]
hp.text = "TRUSTLENS AI  |  RESEARCH REPORT"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    set_font(run, size=8.5, color=GRAY, bold=True)
footer = sec.footer
fp = footer.paragraphs[0]
add_page_number(fp)
for run in fp.runs:
    set_font(run, size=8.5, color=GRAY)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(82)
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TRUSTLENS AI")
set_font(r, size=12, color=BLUE, bold=True)
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Human-Governed Machine Learning\nfor Reliability-Aware Risk Classification")
p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("A reproducible research prototype combining cost-sensitive learning, calibration, drift detection, out-of-distribution screening, explainability and human review")

doc.add_paragraph().paragraph_format.space_after = Pt(24)
meta = doc.add_table(rows=4, cols=2)
set_table_widths(meta, [1.55, 4.95])
meta_data = [
    ("Author", "Akshith Moharampudi"),
    ("Affiliation", "Independent research portfolio; MComp Computer Science, Middlesex University London"),
    ("Project", "github.com/akshith2001/trustlens-ai"),
    ("Version", "Research prototype v0.1 | 22 August 2026"),
]
for i, (label, value) in enumerate(meta_data):
    meta.cell(i, 0).text = label
    meta.cell(i, 1).text = value
    shade(meta.cell(i, 0), PALE)
    for run in meta.cell(i, 0).paragraphs[0].runs:
        set_font(run, bold=True, color=NAVY)
    for run in meta.cell(i, 1).paragraphs[0].runs:
        set_font(run)

doc.add_paragraph().paragraph_format.space_after = Pt(18)
add_callout(doc, "Scope statement.", "This is an educational research prototype. It must not be used to make real lending, employment, legal, medical, insurance or eligibility decisions.", LIGHT)

doc.add_page_break()

add_heading(doc, "Abstract", 1)
add_para(doc, "Machine-learning systems can appear confident even when their probabilities are poorly calibrated, the input population has shifted, or an individual record lies outside the model's supported region. TrustLens AI investigates whether a layered, human-governed auditing architecture can expose these reliability failures more effectively than a prediction score alone. The case study uses the 1,000-record UCI South German Credit dataset and follows a locked development/test protocol. Candidate models were compared using stratified cross-validation, asymmetric error costs, class-specific metrics and calibration quality. The selected configuration was a cost-sensitive random forest with sigmoid calibration, two excluded sensitive or ambiguous fields, and a development-selected decision threshold of 0.20. On the one-time 200-record test partition, the model achieved 0.656 balanced accuracy, 0.833 higher-risk recall and 0.407 precision, with a weighted error cost of 123 versus 300 for a majority-class baseline. Post-hoc Wilson intervals and error-slice diagnostics quantify uncertainty and show that failures are unevenly distributed across recorded feature groups. The system also implements population drift detection, individual out-of-distribution screening, exploratory subgroup diagnostics, local sensitivity explanations and explicit pause/review rules. Results show that the governance layers provide useful audit evidence, but the historical dataset, low precision, subgroup uncertainty and absence of external validation prevent operational claims. The project therefore contributes a reproducible prototype and a transparent account of its limits rather than a deployable lending system.")

add_heading(doc, "Keywords", 2)
add_para(doc, "Responsible AI; machine-learning governance; probability calibration; cost-sensitive classification; drift detection; out-of-distribution detection; explainability; human-in-the-loop review.")

add_heading(doc, "1. Introduction", 1)
add_para(doc, "A classification score is not a complete decision. Its meaning depends on whether probabilities are calibrated, whether current inputs resemble the data used for development, whether errors have asymmetric consequences and whether escalation rules exist when reliability checks fail. These concerns are especially important in high-stakes domains, where an apparently precise output can conceal uncertainty, sampling limitations or unsupported generalisation.")
add_para(doc, "TrustLens AI was developed as a research portfolio project to study this gap between prediction and governance. It places a deterministic governance layer around a tabular classifier and treats monitoring, review and abstention as first-class outputs. The historical credit-risk task is deliberately framed as an educational case study, not as a present-day lending application.")
add_callout(doc, "Research question.", "Can a human-governed auditing layer combining cost-sensitive evaluation, calibration, drift detection, out-of-distribution screening and explainability identify unreliable machine-learning outputs more effectively than confidence scores alone?")

add_heading(doc, "1.1 Contributions", 2)
for item in (
    "A reproducible development protocol that isolates a stratified 20% final test set before model selection.",
    "Joint reporting of discrimination, calibration, asymmetric error cost and class-specific performance.",
    "Locked governance rules that give drift and individual out-of-distribution signals precedence over model confidence.",
    "Exploratory fairness, feature-ablation and explanation evidence with explicit limits on interpretation.",
    "A one-time final evaluation guard, automated tests, public documentation and an interactive research dashboard.",
    "Post-hoc confidence intervals and supported error-slice diagnostics that report uncertainty without retuning the locked model.",
):
    add_bullet(doc, item)

add_heading(doc, "2. Data and ethical framing", 1)
add_para(doc, "The case study uses the UCI South German Credit dataset, which contains 1,000 records and 20 input variables. The normalised target is 0 for lower risk and 1 for higher risk, with 700 and 300 records respectively. The source records were collected in 1973-1975; higher-risk cases were deliberately oversampled, and the transformation applied to credit amount is not fully documented. These properties materially restrict generalisation.")

data_tbl = doc.add_table(rows=1, cols=3)
data_tbl.style = "Table Grid"
for i, h in enumerate(("Data property", "Value", "Research implication")):
    data_tbl.cell(0, i).text = h
for row in (
    ("Records / predictors", "1,000 / 20", "Small sample; uncertainty must be reported"),
    ("Class distribution", "700 lower / 300 higher risk", "Class imbalance motivates class-specific metrics"),
    ("Collection period", "1973-1975", "Not representative of modern lending populations"),
    ("Sensitive coding", "Personal status and sex combined", "Gender fairness is not assessable"),
    ("Intended use", "Education and research", "No real applicant decisions"),
):
    cells = data_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(data_tbl, [1.55, 1.55, 3.4])

add_para(doc, "The project excludes personal_status_sex because its combined encoding prevents a defensible gender interpretation, and excludes foreign_worker from prediction. Age is retained for the governed experiment but audited through broad age bands. Removal of selected features cannot establish fairness, and the available data do not capture all legally or socially relevant groups.")

add_heading(doc, "3. Methodology", 1)
add_heading(doc, "3.1 Experimental protocol", 2)
for item in (
    "Reserve a stratified 20% test partition before model development.",
    "Use stratified five-fold cross-validation on the remaining 80% for candidate comparison.",
    "Choose calibration, feature exclusions and threshold using development evidence only.",
    "Evaluate the locked configuration once on the untouched 200-record partition.",
    "Persist the final metrics to a protected JSON artifact that the evaluation script refuses to overwrite.",
):
    add_bullet(doc, item)

add_heading(doc, "3.2 Models and evaluation", 2)
add_para(doc, "The candidates were a majority-class baseline, logistic regression, cost-sensitive logistic regression, cost-sensitive random forest and cost-sensitive histogram gradient boosting. The principal metrics were balanced accuracy, higher-risk precision, recall and F1, together with Brier score, log loss and expected calibration error. A predefined experimental cost function assigned five units to a false negative and one unit to a false positive. This ratio explores asymmetric consequences; it is not presented as a stakeholder-approved valuation.")

model_tbl = doc.add_table(rows=1, cols=6)
model_tbl.style = "Table Grid"
for i, h in enumerate(("Development model", "Bal. acc.", "Precision", "Recall", "F1", "Cost")):
    model_tbl.cell(0, i).text = h
for row in (
    ("Logistic regression", "0.671", "0.595", "0.487", "0.533", "139.4"),
    ("Cost-sensitive logistic", "0.695", "0.448", "0.825", "0.580", "90.8"),
    ("Cost-sensitive random forest", "0.687", "0.425", "0.896", "0.576", "83.4"),
    ("Cost-sensitive gradient boosting", "0.699", "0.504", "0.692", "0.582", "107.0"),
):
    cells = model_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(model_tbl, [2.25, .85, .85, .75, .75, 1.05])
add_para(doc, "Table 1. Five-fold development-set comparison. Values are means across folds; the final model was selected by considering recall, asymmetric cost, calibration and governance requirements together rather than maximising one metric.", italic=True, after=10)

add_heading(doc, "3.3 Calibration and threshold", 2)
add_para(doc, "The governed random forest uses sigmoid probability calibration with three internal folds. On development data, calibration reduced Brier score from 0.239 to 0.163 and expected calibration error from 0.267 to 0.032. A threshold of 0.20 was selected to target high recall under the predefined five-to-one error-cost ratio. Records above the threshold require human review; near-threshold cases are also escalated.")

add_heading(doc, "3.4 Reliability and governance layers", 2)
gov_tbl = doc.add_table(rows=1, cols=3)
gov_tbl.style = "Table Grid"
for i, h in enumerate(("Signal", "Locked condition", "System response")):
    gov_tbl.cell(0, i).text = h
for row in (
    ("Population drift", "Adversarial AUC >= 0.70", "Pause system and investigate"),
    ("Individual OOD", "Record flagged outside support", "Pause for human review"),
    ("Threshold uncertainty", "Probability within +/-0.05", "Human review required"),
    ("Higher-risk warning", "Probability >= 0.20", "Human review required"),
    ("No trigger", "All checks pass", "Continue monitoring and log evidence"),
):
    cells = gov_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(gov_tbl, [1.55, 2.35, 2.6])
add_para(doc, "Population drift takes precedence over record-level signals, and OOD status takes precedence over predicted probability. A confident-looking score therefore cannot override a detected reliability failure.")

add_heading(doc, "4. Results", 1)
add_heading(doc, "4.1 Locked final evaluation", 2)
add_para(doc, "The selected restricted cost-sensitive random forest was evaluated once on 200 held-out records: 140 lower-risk and 60 higher-risk examples. The decision threshold remained fixed at 0.20.")

res_tbl = doc.add_table(rows=1, cols=4)
res_tbl.style = "Table Grid"
for i, h in enumerate(("Metric", "Result", "Confusion count", "Result")):
    res_tbl.cell(0, i).text = h
for row in (
    ("Accuracy", "0.585", "True negatives", "67"),
    ("Balanced accuracy", "0.656", "False positives", "73"),
    ("Higher-risk precision", "0.407", "False negatives", "10"),
    ("Higher-risk recall", "0.833", "True positives", "50"),
    ("Higher-risk F1", "0.546", "Weighted error cost", "123"),
    ("Brier score", "0.173", "Expected calibration error", "0.070"),
):
    cells = res_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(res_tbl, [1.75, 1.0, 2.5, 1.25])

if FIGURE.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIGURE), width=Inches(5.8))
    add_para(doc, "Figure 1. Locked final-test summary generated from results/final_test_metrics.json.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

add_para(doc, "The majority-class baseline produced 60 false negatives and a weighted error cost of 300. TrustLens reduced this predefined cost by 59%, but generated 73 false positives and a precision of 0.407. These outcomes support the choice to treat the output as a warning for review rather than an automated decision.")

add_heading(doc, "4.2 Statistical uncertainty", 2)
add_para(doc, "Point estimates from a 200-record holdout must not be treated as exact population performance. Two-sided 95% Wilson score intervals were calculated post hoc from the locked confusion matrix. They were not used to select, tune or change the model.")
unc_tbl = doc.add_table(rows=1, cols=3)
unc_tbl.style = "Table Grid"
for i, h in enumerate(("Metric", "Estimate", "95% confidence interval")):
    unc_tbl.cell(0, i).text = h
for row in (
    ("Higher-risk recall", "0.833", "0.720-0.907"),
    ("Higher-risk precision", "0.407", "0.324-0.495"),
    ("Lower-risk specificity", "0.479", "0.398-0.561"),
):
    cells = unc_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(unc_tbl, [2.4, 1.4, 2.7])

add_heading(doc, "4.3 Generalisation gap", 2)
add_para(doc, "Performance weakened between development and the locked test set. Higher-risk recall fell from 0.879 to 0.833, balanced accuracy from 0.698 to 0.656, and expected calibration error rose from 0.032 to 0.070. The project reports this gap as evidence of sampling uncertainty and does not retune against the final partition.")

add_heading(doc, "4.4 Post-hoc error analysis", 2)
add_para(doc, "Descriptive slice analysis was performed after the final evaluation. A slice was reported only when it contained at least ten records from the class relevant to the error rate. The highest observed false-negative rate was associated with checking_account_status=4 (8/11, 0.727), while the highest false-positive rate was associated with checking_account_status=1 (30/32, 0.938). Elevated rates also appeared in several property, job, credit-history, housing, purpose, age and duration slices.")
add_para(doc, "The findings indicate uneven reliability and reinforce the need for human review. They do not establish causation, unfairness or present-day behaviour: the data are historical, individual slices remain small and many comparisons were inspected. The findings were not used to retune the locked model.")

add_heading(doc, "4.5 Drift, OOD and review-budget experiments", 2)
rob_tbl = doc.add_table(rows=1, cols=3)
rob_tbl.style = "Table Grid"
for i, h in enumerate(("Experiment", "Observed result", "Interpretation")):
    rob_tbl.cell(0, i).text = h
for row in (
    ("Random-holdout drift", "AUC 0.485", "No material distinguishability detected"),
    ("Controlled synthetic shift", "AUC 0.779", "Pause threshold correctly triggered"),
    ("Normal-holdout OOD", "7.5% flagged", "False-alert burden remains material"),
    ("Synthetic extreme OOD", "74.0% flagged", "Detector responds to controlled extremes"),
    ("20% review budget", "23.0% errors captured", "25.8% potential cost reduction if corrected"),
):
    cells = rob_tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_repeat_table_header_and_style(rob_tbl, [1.8, 1.55, 3.15])
add_para(doc, "The shifted and extreme samples are synthetic evaluation fixtures rather than deployment events. They verify code paths and governance responses but do not establish real-world monitoring performance.")

add_heading(doc, "4.6 Explainability and subgroup diagnostics", 2)
add_para(doc, "Permutation importance ranked checking-account status as the most influential feature in the development experiment, followed by duration, savings and purpose. Local explanations use controlled feature perturbations to show sensitivity around one record. These methods describe predictive dependence, not causation, valid recourse or financial advice.")
add_para(doc, "Age-band and recorded foreign-worker-code diagnostics report group size, observed warning rate, recall and false-positive rate with Wilson confidence intervals. Small groups produced wide intervals; the source coding prevents a defensible gender analysis. The results are therefore exploratory audit evidence, not proof of fairness or discrimination.")

add_heading(doc, "5. Discussion", 1)
add_heading(doc, "5.1 What the prototype demonstrates", 2)
add_para(doc, "TrustLens demonstrates that reliability governance can be implemented as executable, testable policy rather than as narrative documentation alone. Calibration changes how probabilities should be interpreted; drift and OOD signals can interrupt otherwise confident predictions; and a review budget makes the operational cost of human oversight visible. The locked test protocol also exposes a generalisation gap that would have been hidden by reporting development results only.")

add_heading(doc, "5.2 What the prototype does not demonstrate", 2)
for item in (
    "It does not establish suitability for modern lending populations or any other real decision domain.",
    "It does not prove that the five-to-one error-cost ratio represents affected stakeholders.",
    "It does not establish fairness, causation, recourse quality or reviewer effectiveness.",
    "It does not validate drift or OOD monitoring under live deployment conditions.",
    "It does not remove the substantial false-positive burden created by the high-recall threshold.",
):
    add_bullet(doc, item)

add_heading(doc, "5.3 Relevance to doctoral research", 2)
add_para(doc, "The project motivates research on governance-aware model selection: choosing models not only for predictive performance but also for calibration, abstention quality, monitoring reliability, subgroup uncertainty and the practical limits of human review. A doctoral extension could examine these questions on contemporary, multi-institutional data with stakeholder-defined error costs and prospective evaluation.")

add_heading(doc, "6. Limitations and threats to validity", 1)
limits = (
    ("External validity", "The dataset is historical, small and non-representative; no external cohort is available."),
    ("Construct validity", "The binary target and asymmetric cost function simplify a socially complex decision process."),
    ("Statistical uncertainty", "The 200-record final test produces wide confidence intervals and limits precision."),
    ("Multiple slice inspection", "Ranked error slices are descriptive and may surface unstable chance patterns."),
    ("Fairness measurement", "Gender is not assessable, relevant attributes may be missing, and parity metrics cannot establish justice or non-discrimination."),
    ("Monitoring validity", "Synthetic shifts and extremes test implementation behaviour rather than deployment reliability."),
    ("Human factors", "The review-policy analysis assumes errors presented for review can be corrected; no reviewer study has tested this assumption."),
)
lim_tbl = doc.add_table(rows=1, cols=2)
lim_tbl.style = "Table Grid"
lim_tbl.cell(0, 0).text = "Threat"
lim_tbl.cell(0, 1).text = "Consequence"
for a, b in limits:
    cells = lim_tbl.add_row().cells
    cells[0].text = a
    cells[1].text = b
set_repeat_table_header_and_style(lim_tbl, [1.55, 4.95])

add_heading(doc, "7. Future research", 1)
for item in (
    "Repeat the protocol on a recent, independently collected dataset and report temporal and institutional transportability.",
    "Elicit error costs and review thresholds from affected stakeholders rather than fixing them solely as modelling choices.",
    "Evaluate selective prediction and abstention methods against the current review-budget policy.",
    "Test drift and OOD detectors using realistic temporal, covariate and label shifts with delayed ground truth.",
    "Conduct a human-subject review study measuring error correction, automation bias, workload and explanation usefulness.",
    "Extend the architecture to a second modality, while preserving locked evaluation and governance precedence.",
):
    add_bullet(doc, item)

add_heading(doc, "8. Conclusion", 1)
add_para(doc, "TrustLens AI provides a reproducible demonstration of human-governed machine learning in which calibrated probabilities, population drift, individual support, explanation evidence and escalation rules are evaluated together. The locked final test shows both value and limitation: high recall and lower predefined error cost were achieved, but precision remained low and the generalisation gap was material. The correct conclusion is therefore not that the model is ready for use, but that responsible evaluation requires visible uncertainty, explicit pause conditions, reproducible evidence and honest boundaries around what the data can support.")

add_heading(doc, "References", 1)
refs = [
    "Breiman, L. (2001). Random Forests. Machine Learning, 45, 5-32. https://doi.org/10.1023/A:1010933404324",
    "Dua, D. and Graff, C. (2019). UCI Machine Learning Repository: South German Credit. https://doi.org/10.24432/C5QG88",
    "Niculescu-Mizil, A. and Caruana, R. (2005). Predicting Good Probabilities with Supervised Learning. Proceedings of ICML 2005. https://doi.org/10.1145/1102351.1102430",
    "Platt, J. C. (1999). Probabilistic Outputs for Support Vector Machines and Comparisons to Regularized Likelihood Methods. In Advances in Large Margin Classifiers.",
    "Ribeiro, M. T., Singh, S. and Guestrin, C. (2016). Why Should I Trust You? Explaining the Predictions of Any Classifier. Proceedings of KDD 2016. https://doi.org/10.1145/2939672.2939778",
    "Sculley, D. et al. (2015). Hidden Technical Debt in Machine Learning Systems. Advances in Neural Information Processing Systems 28.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(ref)

add_heading(doc, "Reproducibility statement", 1)
add_para(doc, "The public repository contains the validated data loader, modelling and evaluation code, locked JSON results, uncertainty and error-analysis artifacts, figure-generation script, Streamlit dashboard, model card and 27 automated tests. GitHub Actions executes the test suite on supported Python versions. Repository: https://github.com/akshith2001/trustlens-ai. Live dashboard: https://trustlens-governance-ai.streamlit.app/.")

doc.core_properties.title = "TrustLens AI: Human-Governed Machine Learning for Reliability-Aware Risk Classification"
doc.core_properties.subject = "Responsible AI research report"
doc.core_properties.author = "Akshith Moharampudi"
doc.core_properties.keywords = "responsible AI, machine learning governance, calibration, drift, OOD, explainability"
doc.save(DOCX_PATH)
print(DOCX_PATH)
